"""
Unit tests for document extractors.
Tests PDF, Word, HTML, and spreadsheet extraction capabilities.
"""
import io
import pytest
from unittest.mock import MagicMock, patch

from app.extractors.base import DocumentExtractor, ExtractedContent
from app.extractors.pdf import PDFExtractor
from app.extractors.docx import DocxExtractor
from app.extractors.html import HTMLExtractor
from app.extractors.spreadsheet import SpreadsheetExtractor
from app.extractors.registry import ExtractorRegistry, get_extractor, get_registry


class TestExtractedContent:
    """Tests for ExtractedContent dataclass."""

    def test_create_content(self):
        """Test creating extracted content."""
        content = ExtractedContent(
            text="Test content",
            metadata={"key": "value"},
            pages=5,
        )
        assert content.text == "Test content"
        assert content.metadata["key"] == "value"
        assert content.pages == 5
        assert content.word_count == 2

    def test_auto_word_count(self):
        """Test automatic word count calculation."""
        content = ExtractedContent(text="one two three four five")
        assert content.word_count == 5

    def test_explicit_word_count(self):
        """Test explicit word count is preserved."""
        content = ExtractedContent(text="one two", word_count=100)
        assert content.word_count == 100


class TestExtractorRegistry:
    """Tests for extractor registry."""

    def test_register_extractor(self):
        """Test registering an extractor."""
        registry = ExtractorRegistry()
        extractor = PDFExtractor()
        registry.register(extractor)

        assert registry.get_extractor("document.pdf") is extractor
        assert ".pdf" in registry.supported_extensions()

    def test_get_extractor_by_extension(self):
        """Test getting extractor by file extension."""
        registry = ExtractorRegistry()
        pdf_extractor = PDFExtractor()
        docx_extractor = DocxExtractor()

        registry.register(pdf_extractor)
        registry.register(docx_extractor)

        assert registry.get_extractor("file.pdf") is pdf_extractor
        assert registry.get_extractor("file.docx") is docx_extractor

    def test_get_extractor_case_insensitive(self):
        """Test extension matching is case insensitive."""
        registry = ExtractorRegistry()
        registry.register(PDFExtractor())

        assert registry.get_extractor("FILE.PDF") is not None
        assert registry.get_extractor("file.Pdf") is not None

    def test_unknown_extension_returns_none(self):
        """Test unknown extension returns None."""
        registry = ExtractorRegistry()
        registry.register(PDFExtractor())

        assert registry.get_extractor("file.xyz") is None

    def test_can_extract(self):
        """Test can_extract method."""
        registry = ExtractorRegistry()
        registry.register(PDFExtractor())

        assert registry.can_extract("document.pdf") is True
        assert registry.can_extract("document.xyz") is False

    def test_default_registry(self):
        """Test default registry has all extractors."""
        registry = get_registry()

        assert registry.can_extract("doc.pdf")
        assert registry.can_extract("doc.docx")
        assert registry.can_extract("doc.html")
        assert registry.can_extract("doc.xlsx")
        assert registry.can_extract("doc.csv")

    def test_get_extractor_function(self):
        """Test get_extractor convenience function."""
        assert get_extractor("document.pdf") is not None
        assert get_extractor("document.docx") is not None
        assert get_extractor("document.unknown") is None


class TestPDFExtractor:
    """Tests for PDF document extractor."""

    def test_supported_extensions(self):
        """Test supported extensions."""
        extractor = PDFExtractor()
        assert ".pdf" in extractor.supported_extensions

    def test_can_extract(self):
        """Test can_extract method."""
        extractor = PDFExtractor()
        assert extractor.can_extract("document.pdf") is True
        assert extractor.can_extract("document.PDF") is True
        assert extractor.can_extract("document.txt") is False

    @pytest.mark.asyncio
    async def test_extract_pdf(self):
        """Test extracting text from PDF."""
        extractor = PDFExtractor()

        # Create a mock PDF using pypdf
        from pypdf import PdfWriter

        writer = PdfWriter()
        # Add a blank page - pypdf doesn't easily add text without reportlab
        writer.add_blank_page(width=612, height=792)

        pdf_bytes = io.BytesIO()
        writer.write(pdf_bytes)
        pdf_bytes.seek(0)

        result = await extractor.extract(pdf_bytes, "test.pdf")

        assert isinstance(result, ExtractedContent)
        assert result.metadata["filename"] == "test.pdf"
        assert result.metadata["format"] == "pdf"
        assert result.pages == 1

    @pytest.mark.asyncio
    async def test_extract_pdf_metadata(self):
        """Test extracting PDF metadata."""
        extractor = PDFExtractor()

        # Create PDF with metadata
        from pypdf import PdfWriter

        writer = PdfWriter()
        writer.add_blank_page(width=612, height=792)
        writer.add_metadata({
            "/Title": "Test Document",
            "/Author": "Test Author",
        })

        pdf_bytes = io.BytesIO()
        writer.write(pdf_bytes)
        pdf_bytes.seek(0)

        result = await extractor.extract(pdf_bytes, "test.pdf")

        assert result.metadata.get("title") == "Test Document"
        assert result.metadata.get("author") == "Test Author"


class TestDocxExtractor:
    """Tests for Word document extractor."""

    def test_supported_extensions(self):
        """Test supported extensions."""
        extractor = DocxExtractor()
        assert ".docx" in extractor.supported_extensions

    def test_can_extract(self):
        """Test can_extract method."""
        extractor = DocxExtractor()
        assert extractor.can_extract("document.docx") is True
        assert extractor.can_extract("document.DOCX") is True
        assert extractor.can_extract("document.doc") is False

    @pytest.mark.asyncio
    async def test_extract_docx(self):
        """Test extracting text from Word document."""
        extractor = DocxExtractor()

        # Create a test docx
        from docx import Document

        doc = Document()
        doc.add_paragraph("First paragraph.")
        doc.add_paragraph("Second paragraph with more text.")

        docx_bytes = io.BytesIO()
        doc.save(docx_bytes)
        docx_bytes.seek(0)

        result = await extractor.extract(docx_bytes, "test.docx")

        assert isinstance(result, ExtractedContent)
        assert "First paragraph" in result.text
        assert "Second paragraph" in result.text
        assert result.metadata["filename"] == "test.docx"
        assert result.metadata["format"] == "docx"

    @pytest.mark.asyncio
    async def test_extract_docx_with_tables(self):
        """Test extracting tables from Word document."""
        extractor = DocxExtractor()

        from docx import Document

        doc = Document()
        doc.add_paragraph("Introduction")

        # Add a table
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "Header 1"
        table.cell(0, 1).text = "Header 2"
        table.cell(1, 0).text = "Value 1"
        table.cell(1, 1).text = "Value 2"

        docx_bytes = io.BytesIO()
        doc.save(docx_bytes)
        docx_bytes.seek(0)

        result = await extractor.extract(docx_bytes, "test.docx")

        assert "Header 1" in result.text
        assert "Value 1" in result.text

    @pytest.mark.asyncio
    async def test_extract_docx_metadata(self):
        """Test extracting Word document metadata."""
        extractor = DocxExtractor()

        from docx import Document

        doc = Document()
        doc.core_properties.title = "Test Title"
        doc.core_properties.author = "Test Author"
        doc.add_paragraph("Content")

        docx_bytes = io.BytesIO()
        doc.save(docx_bytes)
        docx_bytes.seek(0)

        result = await extractor.extract(docx_bytes, "test.docx")

        assert result.metadata.get("title") == "Test Title"
        assert result.metadata.get("author") == "Test Author"


class TestHTMLExtractor:
    """Tests for HTML document extractor."""

    def test_supported_extensions(self):
        """Test supported extensions."""
        extractor = HTMLExtractor()
        assert ".html" in extractor.supported_extensions
        assert ".htm" in extractor.supported_extensions
        assert ".xhtml" in extractor.supported_extensions

    def test_can_extract(self):
        """Test can_extract method."""
        extractor = HTMLExtractor()
        assert extractor.can_extract("page.html") is True
        assert extractor.can_extract("page.htm") is True
        assert extractor.can_extract("page.txt") is False

    @pytest.mark.asyncio
    async def test_extract_html(self):
        """Test extracting text from HTML."""
        extractor = HTMLExtractor()

        html_content = b"""
        <!DOCTYPE html>
        <html>
        <head><title>Test Page</title></head>
        <body>
            <h1>Main Heading</h1>
            <p>First paragraph with content.</p>
            <p>Second paragraph.</p>
        </body>
        </html>
        """

        result = await extractor.extract(io.BytesIO(html_content), "test.html")

        assert isinstance(result, ExtractedContent)
        assert "Main Heading" in result.text
        assert "First paragraph" in result.text
        assert result.metadata["title"] == "Test Page"
        assert result.metadata["format"] == "html"

    @pytest.mark.asyncio
    async def test_extract_html_removes_scripts(self):
        """Test that script tags are removed."""
        extractor = HTMLExtractor()

        html_content = b"""
        <html>
        <body>
            <p>Real content</p>
            <script>alert('dangerous');</script>
            <style>.hidden { display: none; }</style>
        </body>
        </html>
        """

        result = await extractor.extract(io.BytesIO(html_content), "test.html")

        assert "Real content" in result.text
        assert "alert" not in result.text
        assert "display: none" not in result.text

    @pytest.mark.asyncio
    async def test_extract_html_removes_navigation(self):
        """Test that nav/footer/header are removed."""
        extractor = HTMLExtractor()

        html_content = b"""
        <html>
        <body>
            <nav>Navigation links</nav>
            <header>Site header</header>
            <main>Main content here</main>
            <footer>Footer content</footer>
        </body>
        </html>
        """

        result = await extractor.extract(io.BytesIO(html_content), "test.html")

        assert "Main content here" in result.text
        assert "Navigation links" not in result.text
        assert "Footer content" not in result.text

    @pytest.mark.asyncio
    async def test_extract_html_meta_description(self):
        """Test extracting meta description."""
        extractor = HTMLExtractor()

        html_content = b"""
        <html>
        <head>
            <title>Page Title</title>
            <meta name="description" content="Page description text">
        </head>
        <body><p>Content</p></body>
        </html>
        """

        result = await extractor.extract(io.BytesIO(html_content), "test.html")

        assert result.metadata.get("description") == "Page description text"


class TestSpreadsheetExtractor:
    """Tests for spreadsheet document extractor."""

    def test_supported_extensions(self):
        """Test supported extensions."""
        extractor = SpreadsheetExtractor()
        assert ".csv" in extractor.supported_extensions
        assert ".xlsx" in extractor.supported_extensions
        assert ".xls" in extractor.supported_extensions

    def test_can_extract(self):
        """Test can_extract method."""
        extractor = SpreadsheetExtractor()
        assert extractor.can_extract("data.csv") is True
        assert extractor.can_extract("data.xlsx") is True
        assert extractor.can_extract("data.txt") is False

    @pytest.mark.asyncio
    async def test_extract_csv(self):
        """Test extracting text from CSV."""
        extractor = SpreadsheetExtractor()

        csv_content = b"Name,Age,City\nAlice,30,NYC\nBob,25,LA\n"

        result = await extractor.extract(io.BytesIO(csv_content), "data.csv")

        assert isinstance(result, ExtractedContent)
        assert "Name" in result.text
        assert "Alice" in result.text
        assert "30" in result.text
        assert result.metadata["format"] == "csv"
        assert result.metadata["row_count"] == 2
        assert result.metadata["column_count"] == 3

    @pytest.mark.asyncio
    async def test_extract_xlsx(self):
        """Test extracting text from Excel."""
        extractor = SpreadsheetExtractor()

        # Create test Excel file
        import pandas as pd

        df = pd.DataFrame({
            "Product": ["Widget", "Gadget"],
            "Price": [10.99, 24.99],
            "Stock": [100, 50],
        })

        xlsx_bytes = io.BytesIO()
        df.to_excel(xlsx_bytes, index=False, engine="openpyxl")
        xlsx_bytes.seek(0)

        result = await extractor.extract(xlsx_bytes, "data.xlsx")

        assert "Product" in result.text
        assert "Widget" in result.text
        assert "10.99" in result.text
        assert result.metadata["format"] == "xlsx"
        assert result.metadata["row_count"] == 2

    @pytest.mark.asyncio
    async def test_extract_xlsx_multiple_sheets(self):
        """Test extracting from Excel with multiple sheets."""
        extractor = SpreadsheetExtractor()

        import pandas as pd

        xlsx_bytes = io.BytesIO()
        with pd.ExcelWriter(xlsx_bytes, engine="openpyxl") as writer:
            pd.DataFrame({"A": [1, 2]}).to_excel(writer, sheet_name="Sheet1", index=False)
            pd.DataFrame({"B": [3, 4]}).to_excel(writer, sheet_name="Sheet2", index=False)
        xlsx_bytes.seek(0)

        result = await extractor.extract(xlsx_bytes, "multi.xlsx")

        assert "Sheet1" in result.text
        assert "Sheet2" in result.text
        assert result.metadata["sheet_count"] == 2
        assert "Sheet1" in result.metadata["sheet_names"]
        assert "Sheet2" in result.metadata["sheet_names"]

    @pytest.mark.asyncio
    async def test_extract_csv_encoding(self):
        """Test CSV with different encodings."""
        extractor = SpreadsheetExtractor()

        # UTF-8 with special characters
        csv_content = "Name,City\nJosé,São Paulo\n".encode("utf-8")

        result = await extractor.extract(io.BytesIO(csv_content), "data.csv")

        assert "José" in result.text
        assert "São Paulo" in result.text


class TestFileProviderWithExtractors:
    """Integration tests for FileDataSourceProvider with extractors."""

    @pytest.mark.asyncio
    async def test_supported_extensions_includes_extractable(self):
        """Test that file provider supports all extractable formats."""
        from app.providers.datasource.file_provider import FileDataSourceProvider

        provider = FileDataSourceProvider()
        extensions = provider.supported_extensions

        # Check text extensions
        assert ".txt" in extensions
        assert ".md" in extensions
        assert ".json" in extensions

        # Check extractable extensions
        assert ".pdf" in extensions
        assert ".docx" in extensions
        assert ".html" in extensions
        assert ".xlsx" in extensions
        assert ".csv" in extensions
