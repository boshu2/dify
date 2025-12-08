"""
Word document extractor using python-docx.

Extracts text content from .docx files for RAG processing.
"""
import io
from typing import BinaryIO

from app.extractors.base import DocumentExtractor, ExtractedContent


class DocxExtractor(DocumentExtractor):
    """Extract text from Word documents using python-docx."""

    @property
    def supported_extensions(self) -> set[str]:
        return {".docx"}

    async def extract(self, file: BinaryIO, filename: str) -> ExtractedContent:
        """Extract text from Word document."""
        from docx import Document

        # Read file content into BytesIO
        content = file.read()
        docx_file = io.BytesIO(content)

        doc = Document(docx_file)
        paragraphs = []
        metadata = {
            "filename": filename,
            "format": "docx",
        }

        # Extract core properties if available
        if doc.core_properties:
            props = doc.core_properties
            if props.title:
                metadata["title"] = props.title
            if props.author:
                metadata["author"] = props.author
            if props.subject:
                metadata["subject"] = props.subject
            if props.keywords:
                metadata["keywords"] = props.keywords

        # Extract paragraph text
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs.append(text)

        # Extract text from tables
        tables_text = []
        for table in doc.tables:
            rows = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                rows.append(" | ".join(cells))
            if rows:
                tables_text.append("\n".join(rows))

        # Combine all text
        all_text = paragraphs + tables_text
        full_text = "\n\n".join(all_text)

        return ExtractedContent(
            text=full_text,
            metadata=metadata,
            pages=1,  # docx doesn't have page concept
        )
